"""Pre-bid query generator service — AI-powered analysis of tender clauses needing clarification."""
import io
import json
import logging
from datetime import date
from uuid import UUID
from typing import List, Optional

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.tender import Tender
from app.models.compliance_matrix import TechnicalComplianceItem

logger = logging.getLogger(__name__)

# Categories of pre-bid queries
QUERY_CATEGORIES = [
    "ambiguous_specs",
    "contradictory_clauses",
    "restrictive_conditions",
    "missing_timelines",
    "unclear_evaluation",
    "onerous_commercial",
]

PREBID_PROMPT = """You are an expert government tender analyst specialising in Indian public procurement (GFR 2017, CVC guidelines, GeM). Analyse the following tender information and identify clauses or provisions that require pre-bid clarification or amendment requests.

TENDER DETAILS:
- Title: {title}
- Reference Number: {reference_number}
- Issuing Authority: {issuing_authority}

TENDER REQUIREMENTS (extracted JSON):
{requirements_json}

COMPLIANCE MATRIX ITEMS (clauses with compliance issues):
{compliance_items_json}

INSTRUCTIONS:
1. Identify clauses that are ambiguous, contradictory, restrictive, missing timelines, have unclear evaluation criteria, or contain onerous commercial terms.
2. For each issue, draft a professional pre-bid query that a bidder would submit.
3. Focus especially on mandatory clauses that are marked as not_complied or partially_complied — these are the most urgent.
4. Reference specific clause numbers and page references where available.
5. Keep clarification language formal and constructive (request amendments, not complaints).

Return ONLY valid JSON in this exact format:
{{
  "queries": [
    {{
      "sno": 1,
      "tender_reference": "Clause 3.2, Page 15",
      "existing_clause": "The exact existing clause or provision text",
      "clarification_sought": "The clarification or amendment request",
      "response": "",
      "category": "ambiguous_specs|contradictory_clauses|restrictive_conditions|missing_timelines|unclear_evaluation|onerous_commercial",
      "priority": "high|medium|low"
    }}
  ]
}}

The "response" field should be left empty — it will be filled in later when the tendering authority responds.
Generate between 5 and 25 queries depending on the complexity of the tender. Prioritise mandatory non-compliant clauses.
"""


async def generate_prebid_queries(
    db: AsyncSession,
    tender_id: UUID,
    tenant_id: str,
) -> dict:
    """Generate pre-bid queries by analysing tender data and compliance matrix.

    Uses Claude AI to identify clauses needing clarification. Falls back to
    compliance-matrix-based generation if Claude is unavailable.

    Args:
        db: Async database session
        tender_id: UUID of the tender
        tenant_id: Tenant ID for multi-tenancy

    Returns:
        dict with queries list, total count, and by_category breakdown
    """
    # 1. Fetch tender details
    result = await db.execute(
        select(Tender).where(
            Tender.id == tender_id,
            Tender.tenant_id == UUID(tenant_id),
        )
    )
    tender = result.scalars().first()
    if not tender:
        raise ValueError(f"Tender {tender_id} not found")

    # 2. Fetch compliance matrix items — focus on problematic ones
    compliance_result = await db.execute(
        select(TechnicalComplianceItem)
        .where(TechnicalComplianceItem.tender_id == tender_id)
        .order_by(TechnicalComplianceItem.sort_order)
    )
    all_compliance_items = compliance_result.scalars().all()

    # Separate problem items for the prompt
    problem_items = [
        item for item in all_compliance_items
        if item.compliance_status in ("not_complied", "partially_complied", "pending")
    ]

    compliance_items_json = json.dumps([
        {
            "clause_number": item.clause_number,
            "clause_title": item.clause_title,
            "clause_description": item.clause_description or "",
            "section": item.section or "",
            "compliance_status": item.compliance_status,
            "deviation_remarks": item.deviation_remarks or "",
            "page_reference": item.page_reference or "",
            "is_mandatory": item.is_mandatory,
            "is_critical": item.is_critical,
        }
        for item in problem_items
    ], indent=2)

    requirements_json = json.dumps(
        tender.requirements if tender.requirements else {},
        indent=2,
        default=str,
    )

    # 3. Build prompt
    prompt = PREBID_PROMPT.format(
        title=tender.title or "N/A",
        reference_number=tender.reference_number or "N/A",
        issuing_authority=tender.issuing_authority or "N/A",
        requirements_json=requirements_json[:50000],  # Cap to avoid context overflow
        compliance_items_json=compliance_items_json[:50000],
    )

    # 4. Call AI via provider factory
    try:
        from app.services.ai_provider_service import send_message

        response_text, tokens_used, model_used = await send_message(
            db=db,
            tenant_id=tenant_id,
            prompt=prompt,
            feature="prebid",
            max_tokens=8192,
        )

        response_text = response_text.strip()

        # Handle markdown code blocks in response
        if response_text.startswith("```"):
            lines = response_text.split("\n")
            json_lines = []
            in_block = False
            for line in lines:
                if line.strip().startswith("```"):
                    in_block = not in_block
                    continue
                if in_block:
                    json_lines.append(line)
            response_text = "\n".join(json_lines)

        parsed = json.loads(response_text)
        queries = parsed.get("queries", [])

        logger.info(
            f"AI generated {len(queries)} pre-bid queries for tender {tender_id} "
            f"({tokens_used} tokens via {model_used})"
        )

    except Exception as e:
        logger.warning(f"AI unavailable for pre-bid queries, using fallback: {e}")
        queries = _generate_fallback_queries(problem_items, tender)

    # 5. Build category breakdown
    by_category: dict[str, int] = {}
    for q in queries:
        cat = q.get("category", "other")
        by_category[cat] = by_category.get(cat, 0) + 1

    return {
        "queries": queries,
        "total": len(queries),
        "by_category": by_category,
        "tender_id": str(tender_id),
        "tender_title": tender.title,
        "tender_reference": tender.reference_number or "",
    }


def _generate_fallback_queries(
    problem_items: list,
    tender,
) -> list:
    """Generate queries from compliance matrix when Claude is unavailable.

    Focuses on mandatory items that are not_complied or partially_complied.
    """
    queries = []
    sno = 0

    for item in problem_items:
        if not item.is_mandatory:
            continue
        if item.compliance_status not in ("not_complied", "partially_complied"):
            continue

        sno += 1

        # Determine category heuristic
        clause_lower = (item.clause_title or "").lower() + " " + (item.clause_description or "").lower()
        category = "ambiguous_specs"
        if any(kw in clause_lower for kw in ["timeline", "schedule", "date", "duration", "period"]):
            category = "missing_timelines"
        elif any(kw in clause_lower for kw in ["evaluation", "scoring", "marks", "weightage"]):
            category = "unclear_evaluation"
        elif any(kw in clause_lower for kw in ["restrict", "only", "exclusive", "minimum", "mandatory experience"]):
            category = "restrictive_conditions"
        elif any(kw in clause_lower for kw in ["penalty", "ld ", "liquidated", "forfeit", "bank guarantee"]):
            category = "onerous_commercial"

        # Build reference string
        ref_parts = []
        if item.clause_number:
            ref_parts.append(f"Clause {item.clause_number}")
        if item.page_reference:
            ref_parts.append(item.page_reference)
        if item.section:
            ref_parts.append(f"Section: {item.section}")
        tender_reference = ", ".join(ref_parts) if ref_parts else "Refer tender document"

        # Build existing clause text
        existing_clause = item.clause_title or ""
        if item.clause_description:
            existing_clause += f" — {item.clause_description}"

        # Build clarification text
        if item.compliance_status == "not_complied":
            clarification = (
                f"The requirement under {item.clause_title} appears restrictive and may limit competition. "
                f"Kindly clarify whether this condition can be relaxed or an equivalent alternative accepted."
            )
            if item.deviation_remarks:
                clarification += f" Bidder's observation: {item.deviation_remarks}"
            priority = "high"
        else:
            clarification = (
                f"The clause regarding {item.clause_title} requires further clarification for accurate compliance. "
                f"Kindly provide additional details or amend to remove ambiguity."
            )
            if item.deviation_remarks:
                clarification += f" Bidder's observation: {item.deviation_remarks}"
            priority = "medium"

        queries.append({
            "sno": sno,
            "tender_reference": tender_reference,
            "existing_clause": existing_clause[:1000],
            "clarification_sought": clarification[:1000],
            "response": "",
            "category": category,
            "priority": priority,
        })

    return queries


SUGGEST_RESPONSE_PROMPT = """You are a senior tendering authority official responding to a pre-bid query from a prospective bidder. Draft a professional, clear response.

TENDER DETAILS:
- Title: {title}
- Reference Number: {reference_number}
- Issuing Authority: {issuing_authority}

PRE-BID QUERY:
- Tender Reference/Clause: {tender_reference}
- Existing Clause: {existing_clause}
- Clarification Sought: {clarification_sought}
- Category: {category}

OTHER QUERIES IN CONTEXT (for consistency):
{other_queries_summary}

INSTRUCTIONS:
1. Draft a response as if you are the tendering authority.
2. Be specific and actionable — clarify, amend, or confirm the clause.
3. Use formal language appropriate for government procurement.
4. If the query requests relaxation, provide a reasonable amendment or explain why the clause stands.
5. Keep the response concise (2-4 sentences).
6. Do NOT use markdown formatting.

Draft the response:"""


async def suggest_response(
    db: AsyncSession,
    tender_id: UUID,
    tenant_id: str,
    query: dict,
    all_queries: list[dict] | None = None,
) -> dict:
    """AI-generate a suggested response for a single pre-bid query.

    Args:
        db: Async database session
        tender_id: UUID of the tender
        tenant_id: Tenant ID
        query: The specific query dict
        all_queries: All queries for context

    Returns:
        dict with suggested_response, tokens_used, model_used
    """
    # Fetch tender for context
    result = await db.execute(
        select(Tender).where(
            Tender.id == tender_id,
            Tender.tenant_id == UUID(tenant_id),
        )
    )
    tender = result.scalars().first()
    if not tender:
        raise ValueError(f"Tender {tender_id} not found")

    # Build context summary from other queries
    other_summary = ""
    if all_queries:
        summaries = []
        for q in all_queries[:10]:
            summaries.append(f"  - [{q.get('category', '')}] {q.get('clarification_sought', '')[:120]}")
        other_summary = "\n".join(summaries)
    else:
        other_summary = "No other queries provided."

    prompt = SUGGEST_RESPONSE_PROMPT.format(
        title=tender.title or "N/A",
        reference_number=tender.reference_number or "N/A",
        issuing_authority=tender.issuing_authority or "N/A",
        tender_reference=query.get("tender_reference", "N/A"),
        existing_clause=query.get("existing_clause", "N/A"),
        clarification_sought=query.get("clarification_sought", "N/A"),
        category=query.get("category", "N/A"),
        other_queries_summary=other_summary,
    )

    from app.services.ai_provider_service import send_message

    response_text, tokens_used, model_used = await send_message(
        db=db,
        tenant_id=tenant_id,
        prompt=prompt,
        feature="prebid",
        max_tokens=1024,
        temperature=0.3,
    )

    return {
        "suggested_response": response_text.strip(),
        "tokens_used": tokens_used,
        "model_used": model_used,
    }


def export_prebid_queries_docx(
    queries: list,
    tender_title: str = "",
    tender_reference: str = "",
) -> io.BytesIO:
    """Generate a DOCX document with pre-bid queries in standard tabular format.

    Args:
        queries: List of query dicts (sno, tender_reference, existing_clause, clarification_sought)
        tender_title: Tender title for the header
        tender_reference: Tender reference number for the header

    Returns:
        BytesIO buffer containing the DOCX file
    """
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = DocxDocument()

    # -- Page margins (narrower for wide table) --
    for section in doc.sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(1.5)

    # -- Title --
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("PRE-BID QUERIES")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

    # -- Subtitle (tender title + reference + date) --
    if tender_title:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(tender_title)
        sub_run.font.size = Pt(13)
        sub_run.bold = True

    meta_parts = []
    if tender_reference:
        meta_parts.append(f"Ref: {tender_reference}")
    meta_parts.append(f"Date: {date.today().strftime('%d %B %Y')}")

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(" | ".join(meta_parts))
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacer

    # -- Table --
    if queries:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        header_cells = table.rows[0].cells
        headers = [
            "S.No",
            "Tender Reference\n(Clause/Page)",
            "Existing Clause/Provision",
            "Clarification/Amendment Sought",
            "Response from\nTendering Authority",
        ]
        for idx, (cell, header_text) in enumerate(zip(header_cells, headers)):
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(header_text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            # Header cell shading
            from docx.oxml.ns import qn
            shading = cell._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(
                qn("w:shd"),
                {
                    qn("w:val"): "clear",
                    qn("w:color"): "auto",
                    qn("w:fill"): "1565C0",
                },
            )
            shading.append(shading_elm)

        # Data rows
        for query in queries:
            row_cells = table.add_row().cells

            # S.No
            row_cells[0].text = str(query.get("sno", ""))
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Tender Reference
            row_cells[1].text = query.get("tender_reference", "")

            # Existing Clause
            row_cells[2].text = query.get("existing_clause", "")

            # Clarification Sought
            row_cells[3].text = query.get("clarification_sought", "")

            # Response
            row_cells[4].text = query.get("response", "")

            # Font size for data cells
            for cell in row_cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

        # Set column widths (approximate)
        for row in table.rows:
            row.cells[0].width = Cm(1.0)
            row.cells[1].width = Cm(3.0)
            row.cells[2].width = Cm(4.5)
            row.cells[3].width = Cm(4.5)
            row.cells[4].width = Cm(5.0)
    else:
        no_data = doc.add_paragraph()
        no_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = no_data.add_run("No pre-bid queries generated.")
        run.font.size = Pt(11)
        run.italic = True

    # -- Footer --
    doc.add_paragraph()  # spacer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(f"Total queries: {len(queries)}")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

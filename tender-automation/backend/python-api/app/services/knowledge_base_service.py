"""Knowledge base service — CRUD, search, suggestions, harvesting, dashboard stats, and doc generation."""
import io
import json
from datetime import datetime, date
from uuid import UUID
from typing import Optional, List

from sqlalchemy import select, func, or_, and_, desc, extract
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.knowledge_base import KnowledgeBaseEntry, KBEntryVersion, KBTenderLink
from app.models.compliance_matrix import TechnicalComplianceItem, BoQLineItem
from app.models.company import CompanyProfile, Personnel
from app.models.reference_bundle import ReferenceBundle
from app.models.tender import Tender


# ---------------------------------------------------------------------------
# Section A: CRUD + Search
# ---------------------------------------------------------------------------

async def list_entries(
    db: AsyncSession,
    tenant_id: str,
    category: Optional[str] = None,
    subcategory: Optional[str] = None,
    search: Optional[str] = None,
    is_current: Optional[bool] = None,
    is_archived: bool = False,
    page: int = 1,
    page_size: int = 50,
) -> dict:
    """Paginated list of KB entries with optional filters."""
    query = (
        select(KnowledgeBaseEntry)
        .where(KnowledgeBaseEntry.tenant_id == UUID(tenant_id))
        .where(KnowledgeBaseEntry.is_archived == is_archived)
    )

    if category:
        query = query.where(KnowledgeBaseEntry.category == category)
    if subcategory:
        query = query.where(KnowledgeBaseEntry.subcategory == subcategory)
    if is_current is not None:
        query = query.where(KnowledgeBaseEntry.is_current == is_current)
    if search:
        pattern = f"%{search}%"
        query = query.where(
            or_(
                KnowledgeBaseEntry.title.ilike(pattern),
                KnowledgeBaseEntry.content.ilike(pattern),
            )
        )

    # Count total
    count_query = select(func.count()).select_from(query.subquery())
    total_result = await db.execute(count_query)
    total = total_result.scalar() or 0

    # Paginate
    offset = (page - 1) * page_size
    query = query.order_by(
        desc(KnowledgeBaseEntry.is_verified),
        desc(KnowledgeBaseEntry.usage_count),
        desc(KnowledgeBaseEntry.updated_at),
    ).offset(offset).limit(page_size)

    result = await db.execute(query)
    entries = result.scalars().all()

    return {
        "items": [_entry_to_dict(e) for e in entries],
        "total": total,
        "page": page,
        "page_size": page_size,
        "total_pages": (total + page_size - 1) // page_size if page_size > 0 else 0,
    }


async def get_entry(db: AsyncSession, entry_id: UUID, tenant_id: str) -> Optional[dict]:
    """Get single entry with its versions and tender links."""
    result = await db.execute(
        select(KnowledgeBaseEntry).where(
            KnowledgeBaseEntry.id == entry_id,
            KnowledgeBaseEntry.tenant_id == UUID(tenant_id),
        )
    )
    entry = result.scalars().first()
    if not entry:
        return None

    # Fetch versions
    versions_result = await db.execute(
        select(KBEntryVersion)
        .where(KBEntryVersion.entry_id == entry_id)
        .order_by(desc(KBEntryVersion.version_number))
    )
    versions = versions_result.scalars().all()

    # Fetch tender links
    links_result = await db.execute(
        select(KBTenderLink)
        .where(KBTenderLink.entry_id == entry_id)
        .order_by(desc(KBTenderLink.used_at))
    )
    links = links_result.scalars().all()

    entry_dict = _entry_to_dict(entry)
    entry_dict["versions"] = [
        {
            "id": str(v.id),
            "version_number": v.version_number,
            "content": v.content,
            "change_reason": v.change_reason,
            "changed_by": str(v.changed_by) if v.changed_by else None,
            "created_at": v.created_at.isoformat() if v.created_at else None,
        }
        for v in versions
    ]
    entry_dict["tender_links"] = [
        {
            "id": str(l.id),
            "tender_id": str(l.tender_id),
            "usage_context": l.usage_context,
            "used_at": l.used_at.isoformat() if l.used_at else None,
            "used_by": str(l.used_by) if l.used_by else None,
        }
        for l in links
    ]
    return entry_dict


async def create_entry(
    db: AsyncSession,
    tenant_id: str,
    user_id: str,
    data: dict,
) -> dict:
    """Create a new KB entry with initial version (v1)."""
    entry = KnowledgeBaseEntry(
        tenant_id=UUID(tenant_id),
        title=data["title"],
        content=data["content"],
        content_format=data.get("content_format", "text"),
        category=data["category"],
        subcategory=data.get("subcategory"),
        tags=data.get("tags"),
        keywords=data.get("keywords"),
        source_type=data.get("source_type", "manual"),
        source_tender_id=UUID(data["source_tender_id"]) if data.get("source_tender_id") else None,
        source_description=data.get("source_description"),
        valid_from=_parse_date(data.get("valid_from")),
        valid_until=_parse_date(data.get("valid_until")),
        fiscal_year=data.get("fiscal_year"),
        is_current=data.get("is_current", True),
        confidence_score=data.get("confidence_score"),
        is_verified=data.get("is_verified", False),
        created_by=UUID(user_id) if user_id else None,
    )
    db.add(entry)
    await db.flush()

    # Create initial version
    version = KBEntryVersion(
        entry_id=entry.id,
        version_number=1,
        content=data["content"],
        change_reason="Initial creation",
        changed_by=UUID(user_id) if user_id else None,
    )
    db.add(version)
    await db.flush()

    return _entry_to_dict(entry)


async def update_entry(
    db: AsyncSession,
    entry_id: UUID,
    tenant_id: str,
    user_id: str,
    data: dict,
) -> Optional[dict]:
    """Update an existing KB entry. Creates a new version if content changed."""
    result = await db.execute(
        select(KnowledgeBaseEntry).where(
            KnowledgeBaseEntry.id == entry_id,
            KnowledgeBaseEntry.tenant_id == UUID(tenant_id),
        )
    )
    entry = result.scalars().first()
    if not entry:
        return None

    content_changed = "content" in data and data["content"] != entry.content

    # Update fields
    updatable = [
        "title", "content", "content_format", "category", "subcategory",
        "tags", "keywords", "source_type", "source_description",
        "fiscal_year", "is_current", "confidence_score", "is_verified", "is_archived",
    ]
    for field in updatable:
        if field in data:
            setattr(entry, field, data[field])

    if "source_tender_id" in data:
        entry.source_tender_id = UUID(data["source_tender_id"]) if data["source_tender_id"] else None
    if "valid_from" in data:
        entry.valid_from = _parse_date(data["valid_from"])
    if "valid_until" in data:
        entry.valid_until = _parse_date(data["valid_until"])

    # Create new version if content changed
    if content_changed:
        # Get latest version number
        ver_result = await db.execute(
            select(func.max(KBEntryVersion.version_number))
            .where(KBEntryVersion.entry_id == entry_id)
        )
        max_ver = ver_result.scalar() or 0

        version = KBEntryVersion(
            entry_id=entry_id,
            version_number=max_ver + 1,
            content=data["content"],
            change_reason=data.get("change_reason", "Content updated"),
            changed_by=UUID(user_id) if user_id else None,
        )
        db.add(version)

    await db.flush()
    return _entry_to_dict(entry)


async def delete_entry(db: AsyncSession, entry_id: UUID, tenant_id: str) -> bool:
    """Hard delete a KB entry and all its versions/links (cascade)."""
    result = await db.execute(
        select(KnowledgeBaseEntry).where(
            KnowledgeBaseEntry.id == entry_id,
            KnowledgeBaseEntry.tenant_id == UUID(tenant_id),
        )
    )
    entry = result.scalars().first()
    if not entry:
        return False
    await db.delete(entry)
    return True


# ---------------------------------------------------------------------------
# Section B: Suggestion Engine
# ---------------------------------------------------------------------------

async def get_suggestions(
    db: AsyncSession,
    tenant_id: str,
    category: Optional[str] = None,
    subcategory: Optional[str] = None,
    context_keywords: Optional[List[str]] = None,
    tender_id: Optional[str] = None,
    limit: int = 20,
) -> list:
    """Get ranked KB suggestions for a given context."""
    query = (
        select(KnowledgeBaseEntry)
        .where(
            KnowledgeBaseEntry.tenant_id == UUID(tenant_id),
            KnowledgeBaseEntry.is_current == True,
            KnowledgeBaseEntry.is_archived == False,
        )
    )

    if category:
        query = query.where(KnowledgeBaseEntry.category == category)
    if subcategory:
        query = query.where(KnowledgeBaseEntry.subcategory == subcategory)

    # Keyword matching on title + content
    if context_keywords:
        keyword_conditions = []
        for kw in context_keywords:
            pattern = f"%{kw}%"
            keyword_conditions.append(
                or_(
                    KnowledgeBaseEntry.title.ilike(pattern),
                    KnowledgeBaseEntry.content.ilike(pattern),
                )
            )
        if keyword_conditions:
            query = query.where(or_(*keyword_conditions))

    query = query.order_by(
        desc(KnowledgeBaseEntry.is_verified),
        desc(KnowledgeBaseEntry.usage_count),
        desc(KnowledgeBaseEntry.updated_at),
    ).limit(limit)

    result = await db.execute(query)
    entries = result.scalars().all()

    # Enrich with staleness and usage-in-tender flags
    now = date.today()
    suggestions = []
    for entry in entries:
        entry_dict = _entry_to_dict(entry)
        entry_dict["is_stale"] = (
            entry.valid_until is not None and entry.valid_until < now
        )

        # Check if already used in the target tender
        already_used = False
        if tender_id:
            link_result = await db.execute(
                select(func.count()).where(
                    KBTenderLink.entry_id == entry.id,
                    KBTenderLink.tender_id == UUID(tender_id),
                )
            )
            already_used = (link_result.scalar() or 0) > 0
        entry_dict["already_used_in_tender"] = already_used
        suggestions.append(entry_dict)

    return suggestions


async def record_usage(
    db: AsyncSession,
    entry_id: UUID,
    tender_id: str,
    usage_context: Optional[str],
    user_id: str,
) -> dict:
    """Record that a KB entry was used in a tender. Increments usage_count and upserts tender link."""
    # Update entry usage stats
    result = await db.execute(
        select(KnowledgeBaseEntry).where(KnowledgeBaseEntry.id == entry_id)
    )
    entry = result.scalars().first()
    if not entry:
        return {"error": "Entry not found"}

    entry.usage_count = (entry.usage_count or 0) + 1
    entry.last_used_at = datetime.utcnow()
    entry.last_used_tender_id = UUID(tender_id)

    # Upsert tender link
    context = usage_context or "general"
    link_result = await db.execute(
        select(KBTenderLink).where(
            KBTenderLink.entry_id == entry_id,
            KBTenderLink.tender_id == UUID(tender_id),
            KBTenderLink.usage_context == context,
        )
    )
    existing_link = link_result.scalars().first()

    if not existing_link:
        link = KBTenderLink(
            entry_id=entry_id,
            tender_id=UUID(tender_id),
            usage_context=context,
            used_by=UUID(user_id) if user_id else None,
        )
        db.add(link)

    await db.flush()
    return _entry_to_dict(entry)


# ---------------------------------------------------------------------------
# Section C: Harvesting Engine
# ---------------------------------------------------------------------------

async def harvest_from_tender(
    db: AsyncSession,
    tender_id: UUID,
    tenant_id: str,
    user_id: str,
) -> dict:
    """Harvest reusable content from a specific tender's compliance matrix and BoQ."""
    created = 0
    skipped = 0

    # 1. Compliance matrix responses (complied items with remarks > 20 chars)
    compliance_result = await db.execute(
        select(TechnicalComplianceItem).where(
            TechnicalComplianceItem.tender_id == tender_id,
            TechnicalComplianceItem.compliance_status == "complied",
        )
    )
    compliance_items = compliance_result.scalars().all()

    # Get tender title for source description
    tender_result = await db.execute(select(Tender).where(Tender.id == tender_id))
    tender = tender_result.scalars().first()
    tender_title = tender.title if tender else "Unknown Tender"

    for item in compliance_items:
        remarks = item.deviation_remarks or ""
        if len(remarks) < 20:
            continue

        content = f"Clause {item.clause_number}: {item.clause_title}\n\nResponse: {remarks}"
        if await _entry_exists(db, UUID(tenant_id), "technical_response", content[:100]):
            skipped += 1
            continue

        entry_data = {
            "title": f"Compliance: {item.clause_title[:200]}",
            "content": content,
            "category": "technical_response",
            "subcategory": "compliance_matrix_response",
            "source_type": "auto_harvested",
            "source_tender_id": str(tender_id),
            "source_description": f"From tender: {tender_title}",
            "tags": [item.section] if item.section else [],
            "keywords": [item.clause_number],
        }
        await create_entry(db, tenant_id, user_id, entry_data)
        created += 1

    # 2. BoQ pricing data (items with unit_rate > 0)
    boq_result = await db.execute(
        select(BoQLineItem).where(
            BoQLineItem.tender_id == tender_id,
            BoQLineItem.unit_rate > 0,
        )
    )
    boq_items = boq_result.scalars().all()

    for item in boq_items:
        content = f"Item: {item.description}\nUnit: {item.unit}\nUnit Rate: {item.unit_rate}\nQuantity: {item.quantity}"
        if item.make_model:
            content += f"\nMake/Model: {item.make_model}"

        if await _entry_exists(db, UUID(tenant_id), "pricing_data", content[:100]):
            skipped += 1
            continue

        entry_data = {
            "title": f"Unit Rate: {item.description[:200]}",
            "content": content,
            "category": "pricing_data",
            "subcategory": item.category or "general",
            "source_type": "auto_harvested",
            "source_tender_id": str(tender_id),
            "source_description": f"From tender: {tender_title}",
            "tags": [item.category] if item.category else [],
        }
        await create_entry(db, tenant_id, user_id, entry_data)
        created += 1

    # 3. Special conditions from tender.requirements JSON
    if tender and tender.requirements and isinstance(tender.requirements, dict):
        special = tender.requirements.get("special_conditions")
        if special and isinstance(special, list):
            for condition in special:
                cond_text = str(condition) if not isinstance(condition, str) else condition
                if len(cond_text) < 10:
                    continue
                if await _entry_exists(db, UUID(tenant_id), "submission_note", cond_text[:100]):
                    skipped += 1
                    continue

                entry_data = {
                    "title": f"Special Condition: {cond_text[:200]}",
                    "content": cond_text,
                    "category": "submission_note",
                    "subcategory": "special_condition",
                    "source_type": "auto_harvested",
                    "source_tender_id": str(tender_id),
                    "source_description": f"From tender: {tender_title}",
                }
                await create_entry(db, tenant_id, user_id, entry_data)
                created += 1

    return {"created": created, "skipped": skipped, "source": "tender", "tender_id": str(tender_id)}


async def harvest_company_data(
    db: AsyncSession,
    tenant_id: str,
    user_id: str,
) -> dict:
    """Sync company master data (turnover, personnel, reference bundles) into KB entries."""
    created = 0
    skipped = 0

    # 1. Annual turnover by FY
    profile_result = await db.execute(
        select(CompanyProfile).where(CompanyProfile.tenant_id == UUID(tenant_id))
    )
    profile = profile_result.scalars().first()

    if profile and profile.annual_turnover and isinstance(profile.annual_turnover, dict):
        for fy, amount in profile.annual_turnover.items():
            content = f"Annual Turnover FY {fy}: INR {amount:,.0f}" if isinstance(amount, (int, float)) else f"Annual Turnover FY {fy}: {amount}"
            if await _entry_exists(db, UUID(tenant_id), "financial_data", f"Turnover FY {fy}"):
                skipped += 1
                continue

            entry_data = {
                "title": f"Annual Turnover FY {fy}",
                "content": content,
                "category": "financial_data",
                "subcategory": "turnover",
                "source_type": "auto_harvested",
                "source_description": "From Company Profile",
                "fiscal_year": fy,
                "is_verified": True,
                "tags": ["turnover", "financial"],
            }
            await create_entry(db, tenant_id, user_id, entry_data)
            created += 1

    # 2. Personnel bios
    if profile:
        personnel_result = await db.execute(
            select(Personnel).where(Personnel.company_id == profile.id)
        )
        personnel_list = personnel_result.scalars().all()

        for person in personnel_list:
            content_parts = [f"Name: {person.name}"]
            if person.designation:
                content_parts.append(f"Designation: {person.designation}")
            if person.role_in_tender:
                content_parts.append(f"Role: {person.role_in_tender}")
            if person.qualification:
                content_parts.append(f"Qualification: {person.qualification}")
            if person.experience_years:
                content_parts.append(f"Experience: {person.experience_years} years")
            if person.specialization:
                specs = person.specialization if isinstance(person.specialization, list) else [person.specialization]
                content_parts.append(f"Specialization: {', '.join(str(s) for s in specs)}")
            content = "\n".join(content_parts)

            if await _entry_exists(db, UUID(tenant_id), "personnel_bio", person.name):
                skipped += 1
                continue

            entry_data = {
                "title": f"{person.name} — {person.designation or person.role_in_tender or 'Personnel'}",
                "content": content,
                "category": "personnel_bio",
                "subcategory": person.role_in_tender or "general",
                "source_type": "auto_harvested",
                "source_description": "From Company Profile — Personnel",
                "is_verified": True,
                "tags": ["personnel", person.role_in_tender] if person.role_in_tender else ["personnel"],
            }
            await create_entry(db, tenant_id, user_id, entry_data)
            created += 1

    # 3. Experience summaries from reference bundles
    bundles_result = await db.execute(
        select(ReferenceBundle).where(
            ReferenceBundle.tenant_id == UUID(tenant_id),
            ReferenceBundle.status == "completed",
        )
    )
    bundles = bundles_result.scalars().all()

    for bundle in bundles:
        content_parts = [
            f"Client: {bundle.client_name}",
            f"Project: {bundle.project_name or bundle.bundle_name}",
        ]
        if bundle.contract_value:
            content_parts.append(f"Contract Value: INR {bundle.contract_value:,.0f}")
        if bundle.duration_months:
            content_parts.append(f"Duration: {bundle.duration_months} months")
        if bundle.scope_description:
            content_parts.append(f"Scope: {bundle.scope_description}")
        if bundle.work_order_number:
            content_parts.append(f"Work Order: {bundle.work_order_number}")
        content = "\n".join(content_parts)

        if await _entry_exists(db, UUID(tenant_id), "experience_summary", bundle.client_name):
            skipped += 1
            continue

        tags = []
        if bundle.client_type_tags and isinstance(bundle.client_type_tags, list):
            tags.extend(str(t) for t in bundle.client_type_tags)
        if bundle.work_type_tags and isinstance(bundle.work_type_tags, list):
            tags.extend(str(t) for t in bundle.work_type_tags)

        entry_data = {
            "title": f"Experience: {bundle.client_name} — {bundle.project_name or bundle.bundle_name}",
            "content": content,
            "category": "experience_summary",
            "subcategory": bundle.client_type or "general",
            "source_type": "auto_harvested",
            "source_description": f"From Reference Bundle: {bundle.bundle_name}",
            "is_verified": True,
            "tags": tags or ["experience"],
        }
        await create_entry(db, tenant_id, user_id, entry_data)
        created += 1

    return {"created": created, "skipped": skipped, "source": "company"}


async def get_category_stats(db: AsyncSession, tenant_id: str) -> list:
    """Get per-category counts, usage totals, and stale counts."""
    result = await db.execute(
        select(
            KnowledgeBaseEntry.category,
            func.count(KnowledgeBaseEntry.id).label("count"),
            func.sum(KnowledgeBaseEntry.usage_count).label("total_usage"),
        )
        .where(
            KnowledgeBaseEntry.tenant_id == UUID(tenant_id),
            KnowledgeBaseEntry.is_archived == False,
        )
        .group_by(KnowledgeBaseEntry.category)
    )
    rows = result.all()

    now = date.today()
    stats = []
    for row in rows:
        # Count stale entries in this category
        stale_result = await db.execute(
            select(func.count()).where(
                KnowledgeBaseEntry.tenant_id == UUID(tenant_id),
                KnowledgeBaseEntry.category == row.category,
                KnowledgeBaseEntry.is_archived == False,
                KnowledgeBaseEntry.valid_until < now,
            )
        )
        stale_count = stale_result.scalar() or 0

        stats.append({
            "category": row.category,
            "count": row.count,
            "total_usage": row.total_usage or 0,
            "stale_count": stale_count,
        })

    return stats


# ---------------------------------------------------------------------------
# Section D: Dashboard Stats
# ---------------------------------------------------------------------------

async def dashboard_stats(db: AsyncSession, tenant_id: str) -> dict:
    """Return aggregate KB stats for the tenant dashboard."""
    tid = UUID(tenant_id)
    now = date.today()

    # Total non-archived entries
    total_result = await db.execute(
        select(func.count()).where(
            KnowledgeBaseEntry.tenant_id == tid,
            KnowledgeBaseEntry.is_archived == False,
        )
    )
    total_entries = total_result.scalar() or 0

    # Verified count
    verified_result = await db.execute(
        select(func.count()).where(
            KnowledgeBaseEntry.tenant_id == tid,
            KnowledgeBaseEntry.is_archived == False,
            KnowledgeBaseEntry.is_verified == True,
        )
    )
    verified_count = verified_result.scalar() or 0

    # Stale count (valid_until < today)
    stale_result = await db.execute(
        select(func.count()).where(
            KnowledgeBaseEntry.tenant_id == tid,
            KnowledgeBaseEntry.is_archived == False,
            KnowledgeBaseEntry.valid_until < now,
        )
    )
    stale_count = stale_result.scalar() or 0

    # Harvested this month
    month_start = date(now.year, now.month, 1)
    harvested_result = await db.execute(
        select(func.count()).where(
            KnowledgeBaseEntry.tenant_id == tid,
            KnowledgeBaseEntry.source_type == "auto_harvested",
            KnowledgeBaseEntry.created_at >= datetime.combine(month_start, datetime.min.time()),
        )
    )
    harvested_this_month = harvested_result.scalar() or 0

    # Category distribution
    cat_result = await db.execute(
        select(
            KnowledgeBaseEntry.category,
            func.count(KnowledgeBaseEntry.id).label("count"),
        )
        .where(
            KnowledgeBaseEntry.tenant_id == tid,
            KnowledgeBaseEntry.is_archived == False,
        )
        .group_by(KnowledgeBaseEntry.category)
    )
    category_distribution = [
        {"category": row.category, "count": row.count}
        for row in cat_result.all()
    ]

    return {
        "total_entries": total_entries,
        "verified_count": verified_count,
        "stale_count": stale_count,
        "harvested_this_month": harvested_this_month,
        "category_distribution": category_distribution,
    }


# ---------------------------------------------------------------------------
# Section E: Document Generation
# ---------------------------------------------------------------------------

async def generate_summary(
    db: AsyncSession,
    tender_id: UUID,
    entry_ids: List[str],
    tenant_id: str,
) -> io.BytesIO:
    """Generate a DOCX summary document from selected KB entries grouped by category."""
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Fetch tender title
    tender_result = await db.execute(select(Tender).where(Tender.id == tender_id))
    tender = tender_result.scalars().first()
    tender_title = tender.title if tender else "Knowledge Base Summary"

    # Fetch entries
    uuids = [UUID(eid) for eid in entry_ids]
    result = await db.execute(
        select(KnowledgeBaseEntry).where(
            KnowledgeBaseEntry.id.in_(uuids),
            KnowledgeBaseEntry.tenant_id == UUID(tenant_id),
        )
    )
    entries = result.scalars().all()

    # Group by category
    grouped: dict[str, list] = {}
    for entry in entries:
        cat = entry.category or "other"
        if cat not in grouped:
            grouped[cat] = []
        grouped[cat].append(entry)

    # Build DOCX
    doc = DocxDocument()

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("KNOWLEDGE BASE SUMMARY")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(tender_title)
    sub_run.font.size = Pt(14)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated: {date.today().strftime('%d %B %Y')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacer

    category_labels = {
        "company_snippet": "Company Information",
        "financial_data": "Financial Data",
        "experience_summary": "Experience Summaries",
        "declaration_text": "Declaration Texts",
        "technical_response": "Technical Responses",
        "personnel_bio": "Personnel Biographies",
        "pricing_data": "Pricing Data",
        "standard_clause": "Standard Clauses",
        "submission_note": "Submission Notes",
    }

    entry_num = 0
    for category, cat_entries in grouped.items():
        # Category heading
        cat_label = category_labels.get(category, category.replace("_", " ").title())
        heading = doc.add_heading(cat_label, level=1)

        for entry in cat_entries:
            entry_num += 1
            # Entry title
            entry_heading = doc.add_heading(f"{entry_num}. {entry.title}", level=2)

            # Metadata line
            meta_parts = []
            if entry.subcategory:
                meta_parts.append(f"Subcategory: {entry.subcategory}")
            if entry.fiscal_year:
                meta_parts.append(f"FY: {entry.fiscal_year}")
            if entry.is_verified:
                meta_parts.append("Verified")
            if meta_parts:
                meta_para = doc.add_paragraph()
                meta_run = meta_para.add_run(" | ".join(meta_parts))
                meta_run.font.size = Pt(9)
                meta_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                meta_run.italic = True

            # Content
            doc.add_paragraph(entry.content)

            # Separator
            doc.add_paragraph("_" * 60)

    # Summary footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(f"Total entries: {entry_num} | Categories: {len(grouped)}")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

async def _entry_exists(db: AsyncSession, tenant_id: UUID, category: str, content_snippet: str) -> bool:
    """Check if a similar entry already exists (idempotency guard)."""
    result = await db.execute(
        select(func.count()).where(
            KnowledgeBaseEntry.tenant_id == tenant_id,
            KnowledgeBaseEntry.category == category,
            or_(
                KnowledgeBaseEntry.title.ilike(f"%{content_snippet[:100]}%"),
                KnowledgeBaseEntry.content.ilike(f"%{content_snippet[:100]}%"),
            ),
        )
    )
    return (result.scalar() or 0) > 0


def _parse_date(value) -> Optional[date]:
    """Parse a date string (YYYY-MM-DD) or return None."""
    if not value:
        return None
    if isinstance(value, date):
        return value
    try:
        return datetime.strptime(str(value), "%Y-%m-%d").date()
    except (ValueError, TypeError):
        return None


def _entry_to_dict(entry: KnowledgeBaseEntry) -> dict:
    return {
        "id": str(entry.id),
        "tenant_id": str(entry.tenant_id),
        "title": entry.title,
        "content": entry.content,
        "content_format": entry.content_format,
        "category": entry.category,
        "subcategory": entry.subcategory,
        "tags": entry.tags,
        "keywords": entry.keywords,
        "source_type": entry.source_type,
        "source_tender_id": str(entry.source_tender_id) if entry.source_tender_id else None,
        "source_description": entry.source_description,
        "valid_from": entry.valid_from.isoformat() if entry.valid_from else None,
        "valid_until": entry.valid_until.isoformat() if entry.valid_until else None,
        "fiscal_year": entry.fiscal_year,
        "is_current": entry.is_current,
        "usage_count": entry.usage_count,
        "last_used_at": entry.last_used_at.isoformat() if entry.last_used_at else None,
        "last_used_tender_id": str(entry.last_used_tender_id) if entry.last_used_tender_id else None,
        "confidence_score": entry.confidence_score,
        "is_verified": entry.is_verified,
        "is_archived": entry.is_archived,
        "created_by": str(entry.created_by) if entry.created_by else None,
        "created_at": entry.created_at.isoformat() if entry.created_at else None,
        "updated_at": entry.updated_at.isoformat() if entry.updated_at else None,
    }


# ---------------------------------------------------------------------------
# Section F: Auto-Fill Suggestions
# ---------------------------------------------------------------------------

# Category mapping for target areas
_TARGET_CATEGORIES = {
    "declarations": ["declaration_text", "standard_clause"],
    "compliance": ["technical_response", "company_snippet", "experience_summary"],
    "bid_preparation": ["technical_response", "pricing_data", "personnel_bio"],
}


async def auto_fill_suggestions(
    db: AsyncSession,
    tenant_id: str,
    tender_id: str,
    target_area: str,
) -> list[dict]:
    """
    Match KB entries to tender requirements for auto-fill.

    For each target area, fetches relevant KB entries and scores them
    based on verification status, recency, and usage count.
    """
    categories = _TARGET_CATEGORIES.get(target_area, [])
    if not categories:
        return []

    # Fetch verified, current entries in matching categories
    query = (
        select(KnowledgeBaseEntry)
        .where(
            and_(
                KnowledgeBaseEntry.tenant_id == UUID(tenant_id),
                KnowledgeBaseEntry.is_archived == False,
                KnowledgeBaseEntry.category.in_(categories),
            )
        )
        .order_by(
            desc(KnowledgeBaseEntry.is_verified),
            desc(KnowledgeBaseEntry.usage_count),
            desc(KnowledgeBaseEntry.updated_at),
        )
        .limit(30)
    )

    result = await db.execute(query)
    entries = result.scalars().all()

    # Score each entry
    suggestions = []
    for entry in entries:
        confidence = 0.5
        if entry.is_verified:
            confidence += 0.25
        if entry.usage_count and entry.usage_count > 3:
            confidence += 0.15
        if entry.is_current:
            confidence += 0.10

        target_field = _infer_target_field(target_area, entry.category, entry.subcategory)

        suggestions.append({
            "kb_entry_id": str(entry.id),
            "title": entry.title,
            "content": entry.content,
            "target_field": target_field,
            "confidence_score": round(confidence, 2),
            "category": entry.category,
            "subcategory": entry.subcategory,
            "is_verified": entry.is_verified,
            "usage_count": entry.usage_count or 0,
        })

    # Sort by confidence descending
    suggestions.sort(key=lambda s: s["confidence_score"], reverse=True)
    return suggestions[:20]


def _infer_target_field(target_area: str, category: str, subcategory: str | None) -> str:
    """Infer which field/section the entry should auto-fill into."""
    if target_area == "declarations":
        if category == "declaration_text":
            return subcategory or "general_declaration"
        return "standard_clause"
    elif target_area == "compliance":
        if category == "technical_response":
            return "technical_response"
        if category == "experience_summary":
            return "experience_narrative"
        return "company_overview"
    elif target_area == "bid_preparation":
        if category == "pricing_data":
            return "boq_reference"
        if category == "personnel_bio":
            return "team_profile"
        return "technical_approach"
    return "general"

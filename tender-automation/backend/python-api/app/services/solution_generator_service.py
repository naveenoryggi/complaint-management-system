"""Solution document generator — end-to-end AI pipeline for tender response generation.

Pipeline steps:
1. Analyze tender structure (detect sections)
2. Classify requirements into proposal sections
3. Generate dynamic TOC
4. Generate content per section (with requirement traceability)
5. Suggest diagrams
6. Build compliance matrix
7. Run gap analysis
8. Simulate evaluation scoring
9. Optimize weak sections
10. Assemble final DOCX
"""
import json
import logging
from datetime import datetime, timezone
from uuid import UUID

from sqlalchemy import select, delete
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.tender import Tender
from app.models.compliance_matrix import TechnicalComplianceItem
from app.models.company import CompanyProfile
from app.models.reference_bundle import ReferenceBundle, TenderCriteria
from app.models.knowledge_base import KnowledgeBaseEntry
from app.models.proposal import (
    TenderSection,
    ProposalSection,
    GeneratedContent,
    ProposalGeneration,
)

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# System prompt — sent with every AI call
# ---------------------------------------------------------------------------

SYSTEM_PROMPT = """You are a senior solution architect and government tender response specialist.

You prepare high-scoring technical bid responses for Indian government tenders (GeM, CPPP, state portals).

Guidelines:
1. Never invent tender requirements — only reference what is provided.
2. Always maintain traceability to requirement IDs when provided.
3. Use structured proposal writing style with tables, bullet points, and numbered lists.
4. Avoid marketing language — use technical, professional tone.
5. Insert placeholders where bidder-specific input is required: [Bidder Input Required - Description]
6. Insert diagram placeholders where visual elements would help: [Insert Diagram - Description]
7. Outputs must be structured and suitable for automated document assembly.
8. Write in third person ("The bidder proposes..." not "We propose...")."""


# ---------------------------------------------------------------------------
# Step 1: Tender Structure Analysis
# ---------------------------------------------------------------------------

STRUCTURE_ANALYSIS_PROMPT = """Analyze the following tender document details and identify the major sections present.

TENDER TITLE: {title}
REFERENCE: {reference_number}
ISSUING AUTHORITY: {issuing_authority}

EXTRACTED REQUIREMENTS:
{requirements_json}

Identify which of these standard tender sections are present or implied:
- Scope of Work
- Technical Specifications
- Functional Requirements
- Eligibility Criteria
- Evaluation Criteria
- Documentation Requirements
- Implementation Requirements
- Training Requirements
- Support & Maintenance Requirements
- SLA Requirements
- Security Requirements
- Integration Requirements
- Manpower / Resource Requirements
- Hardware / Infrastructure Requirements

Return ONLY valid JSON:
{{
  "sections": [
    {{
      "section_name": "Scope of Work",
      "description": "Brief description of what this section covers",
      "key_requirements": ["List of key requirements from this section"],
      "page_reference": "Page/section reference if available"
    }}
  ]
}}

Only include sections that are actually present or clearly implied in the tender."""


# ---------------------------------------------------------------------------
# Step 2: Requirement Classification
# ---------------------------------------------------------------------------

REQUIREMENT_CLASSIFICATION_PROMPT = """Classify the following tender requirements into proposal response sections.

TENDER: {title}
REQUIREMENTS:
{requirements_list}

Map each requirement to the most appropriate proposal section. A requirement may map to multiple sections.

Standard proposal sections:
1. Executive Summary
2. Bidder Profile & Organization
3. Financial Capability
4. Relevant Experience
5. Understanding of the Project
6. Proposed Solution Architecture
7. System Modules & Features
8. Integration Framework
9. Security Framework
10. Implementation Methodology
11. Project Management Plan
12. Resource Deployment Plan
13. Training Plan
14. Support & Maintenance Plan
15. SLA Compliance
16. Risk Management
17. Transition Plan
18. Exit Management
19. Compliance Matrix

Return ONLY valid JSON:
{{
  "sections": [
    {{
      "section_number": "6",
      "section_title": "Proposed Solution Architecture",
      "mapped_requirements": ["REQ-001", "REQ-005"],
      "importance": "high",
      "estimated_pages": 5
    }}
  ]
}}

Only include sections that are relevant to this specific tender. Skip sections with no mapped requirements."""


# ---------------------------------------------------------------------------
# Step 3: Dynamic TOC Generation
# ---------------------------------------------------------------------------

TOC_GENERATION_PROMPT = """Generate a detailed Table of Contents for the tender response document.

TENDER: {title}
REFERENCE: {reference_number}
CLASSIFIED SECTIONS: {classified_sections}
EVALUATION CRITERIA: {evaluation_criteria}

Rules:
1. Only include sections relevant to THIS tender.
2. Include bidder information sections if eligibility criteria require them.
3. Include lifecycle sections (training, support, exit) only if the tender requires them.
4. Maintain hierarchical numbering (1, 1.1, 1.1.1).
5. Order sections logically: company info first, then technical, then operational.
6. Add sub-sections where the tender has detailed requirements.
7. Ensure every evaluation criteria has at least one corresponding section.

Return ONLY valid JSON:
{{
  "toc": [
    {{
      "section_number": "1",
      "section_title": "Executive Summary",
      "parent_section": null,
      "mapped_requirement_ids": [],
      "estimated_pages": 2
    }},
    {{
      "section_number": "6",
      "section_title": "Proposed Solution Architecture",
      "parent_section": null,
      "mapped_requirement_ids": ["uuid1", "uuid2"],
      "estimated_pages": 5
    }},
    {{
      "section_number": "6.1",
      "section_title": "High-Level Architecture",
      "parent_section": "6",
      "mapped_requirement_ids": ["uuid1"],
      "estimated_pages": 2
    }}
  ]
}}"""


# ---------------------------------------------------------------------------
# Step 4: Section Content Generation
# ---------------------------------------------------------------------------

SECTION_CONTENT_PROMPT = """Generate the following section of the tender response document.

TENDER: {title}
REFERENCE: {reference_number}

SECTION: {section_number} — {section_title}

MAPPED REQUIREMENTS:
{mapped_requirements}

COMPANY PROFILE:
{company_profile}

RELEVANT EXPERIENCE:
{relevant_experience}

KNOWLEDGE BASE CONTENT:
{kb_content}

EVALUATION CRITERIA FOR THIS SECTION:
{evaluation_criteria}

Instructions:
1. Begin with a brief overview paragraph.
2. Address EVERY mapped requirement explicitly with traceability.
3. Use the company profile and experience data to personalize the response.
4. Use knowledge base content where relevant — do not copy verbatim, adapt to context.
5. Include tables for specifications, feature lists, and compliance mapping.
6. Insert diagram placeholders: [Insert Diagram - Description]
7. Mark bidder-specific inputs: [Bidder Input Required - Description]
8. Write {estimated_pages} pages worth of content (approximately {word_count} words).
9. Use formal proposal language, third person.
10. For architecture sections, describe components, data flow, scalability, and security.
11. For methodology sections, describe phases, activities, deliverables, and timelines.

Format: Use markdown with proper headings (##, ###), bullet points, and tables."""


# ---------------------------------------------------------------------------
# Step 5: Diagram Suggestions
# ---------------------------------------------------------------------------

DIAGRAM_SUGGESTION_PROMPT = """Based on the following proposal sections and content, suggest diagrams and visual elements.

TENDER: {title}
SECTIONS:
{sections_summary}

For each suggestion provide:
- Which section it belongs to
- Diagram type (architecture, flow, topology, timeline, org chart, matrix, infographic)
- Description of what the diagram should show
- Why it improves the proposal

Return ONLY valid JSON:
{{
  "diagrams": [
    {{
      "section_number": "6",
      "section_title": "Proposed Solution Architecture",
      "diagram_type": "architecture",
      "description": "High-level system architecture showing all major components",
      "rationale": "Demonstrates technical understanding to evaluators"
    }}
  ]
}}"""


# ---------------------------------------------------------------------------
# Step 6: Gap Analysis
# ---------------------------------------------------------------------------

GAP_ANALYSIS_PROMPT = """Analyze the generated tender response for compliance gaps.

TENDER: {title}
TOTAL REQUIREMENTS: {total_requirements}

REQUIREMENTS LIST:
{requirements_list}

GENERATED SECTIONS:
{sections_with_content}

Check:
1. Are ALL requirements addressed? List any missing ones.
2. Are explanations technically sufficient?
3. Are there inconsistencies between sections?
4. Are mandatory requirements addressed with enough depth?
5. Are there any requirements addressed superficially?

Return ONLY valid JSON:
{{
  "total_requirements": {total_requirements},
  "covered_requirements": 0,
  "coverage_percentage": 0,
  "gaps": [
    {{
      "requirement_id": "uuid or description",
      "issue": "What is missing or weak",
      "risk_level": "critical|high|medium|low",
      "recommendation": "How to fix",
      "affected_section": "Section number"
    }}
  ]
}}"""


# ---------------------------------------------------------------------------
# Step 7: Evaluation Simulation
# ---------------------------------------------------------------------------

EVALUATION_SIMULATION_PROMPT = """Simulate a technical evaluation committee scoring of this tender response.

TENDER: {title}
EVALUATION CRITERIA:
{evaluation_criteria}

PROPOSAL SECTIONS AND CONTENT:
{sections_with_content}

COMPANY PROFILE:
{company_profile}

RELEVANT EXPERIENCE:
{relevant_experience}

For each evaluation criterion:
1. Assess the proposal content quality against what evaluators typically expect.
2. Estimate a realistic score (be conservative, not optimistic).
3. Provide specific improvement suggestions.

Return ONLY valid JSON:
{{
  "evaluation": [
    {{
      "criteria": "Technical Architecture",
      "max_marks": 25,
      "estimated_score": 18,
      "percentage": 72,
      "strengths": ["Good component coverage"],
      "weaknesses": ["Missing scalability details"],
      "improvement_suggestions": ["Add horizontal scaling approach"]
    }}
  ],
  "total_max_marks": 100,
  "total_estimated_score": 0,
  "overall_percentage": 0,
  "summary": "Overall assessment",
  "top_improvement_areas": ["List of 3-5 areas to focus on"]
}}"""


# ---------------------------------------------------------------------------
# Step 8: Section Optimizer
# ---------------------------------------------------------------------------

SECTION_OPTIMIZER_PROMPT = """Improve the following tender response section to maximize technical evaluation score.

SECTION: {section_number} — {section_title}

CURRENT CONTENT:
{current_content}

EVALUATION FEEDBACK:
{evaluation_feedback}

GAP ANALYSIS FINDINGS:
{gap_findings}

MAPPED REQUIREMENTS (must all be addressed):
{mapped_requirements}

Instructions:
1. Enhance technical clarity and depth.
2. Strengthen implementation methodology descriptions.
3. Improve architecture explanations.
4. Ensure all mapped requirements are explicitly addressed with traceability.
5. Add specific technical details where the current content is vague.
6. Maintain formal proposal language.
7. Keep all existing diagram placeholders and bidder input markers.
8. Do NOT add fictional capabilities — only enhance explanations of what is already claimed.

Return the improved section content in markdown format."""


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

async def _ai_call(db, tenant_id, prompt, feature, max_tokens=4096, temperature=0.15):
    """Call AI provider with the system prompt."""
    from app.services.ai_provider_service import send_message

    full_prompt = f"{SYSTEM_PROMPT}\n\n---\n\n{prompt}"

    response_text, tokens_used, model_used = await send_message(
        db=db,
        tenant_id=tenant_id,
        prompt=full_prompt,
        feature=feature,
        max_tokens=max_tokens,
        temperature=temperature,
    )

    return response_text.strip(), tokens_used, model_used


def _parse_json_response(text: str) -> dict:
    """Parse JSON from AI response, stripping markdown code blocks if present."""
    text = text.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        lines = [l for l in lines if not l.strip().startswith("```")]
        text = "\n".join(lines)
    return json.loads(text)


async def _get_company_profile_text(db: AsyncSession, tenant_id: str) -> str:
    """Get company profile as text for prompt context."""
    result = await db.execute(
        select(CompanyProfile).where(CompanyProfile.tenant_id == UUID(tenant_id))
    )
    profile = result.scalars().first()
    if not profile:
        return "Company profile not available. [Bidder Input Required - Company Details]"

    parts = []
    if profile.company_name:
        parts.append(f"Company Name: {profile.company_name}")
    if profile.registered_address:
        parts.append(f"Registered Address: {profile.registered_address}")
    if profile.pan_number:
        parts.append(f"PAN: {profile.pan_number}")
    if profile.gstin:
        parts.append(f"GSTIN: {profile.gstin}")
    if profile.year_of_incorporation:
        parts.append(f"Year of Incorporation: {profile.year_of_incorporation}")
    if profile.employee_count:
        parts.append(f"Employee Count: {profile.employee_count}")
    if profile.annual_turnover:
        parts.append(f"Annual Turnover: {profile.annual_turnover}")
    if profile.net_worth:
        parts.append(f"Net Worth: {profile.net_worth}")
    if profile.about_company:
        parts.append(f"About: {profile.about_company}")
    return "\n".join(parts) if parts else "Minimal company profile available."


async def _get_experience_text(db: AsyncSession, tenant_id: str, limit: int = 10) -> str:
    """Get relevant experience as text."""
    result = await db.execute(
        select(ReferenceBundle)
        .where(ReferenceBundle.tenant_id == UUID(tenant_id))
        .order_by(ReferenceBundle.contract_value.desc().nulls_last())
        .limit(limit)
    )
    bundles = result.scalars().all()
    if not bundles:
        return "No past project experience recorded. [Bidder Input Required - Project Experience]"

    parts = []
    for b in bundles:
        line = f"- {b.bundle_name}: Client={b.client_name}"
        if b.contract_value:
            line += f", Value={b.contract_value}"
        if b.duration_months:
            line += f", Duration={b.duration_months} months"
        if b.scope_description:
            line += f", Scope: {b.scope_description[:200]}"
        parts.append(line)
    return "\n".join(parts)


async def _get_kb_content(db: AsyncSession, tenant_id: str, category: str = None) -> str:
    """Get knowledge base content for prompt context."""
    query = select(KnowledgeBaseEntry).where(
        KnowledgeBaseEntry.tenant_id == UUID(tenant_id),
        KnowledgeBaseEntry.is_current == True,
        KnowledgeBaseEntry.is_archived == False,
    )
    if category:
        query = query.where(KnowledgeBaseEntry.category == category)
    query = query.order_by(KnowledgeBaseEntry.usage_count.desc()).limit(20)

    result = await db.execute(query)
    entries = result.scalars().all()
    if not entries:
        return "No knowledge base entries available."

    parts = []
    for e in entries:
        parts.append(f"[{e.category}] {e.title}: {e.content[:300]}")
    return "\n".join(parts)


async def _get_requirements_list(db: AsyncSession, tender_id: UUID) -> list[dict]:
    """Get compliance items as structured requirement list."""
    result = await db.execute(
        select(TechnicalComplianceItem)
        .where(TechnicalComplianceItem.tender_id == tender_id)
        .order_by(TechnicalComplianceItem.sort_order)
    )
    items = result.scalars().all()
    return [
        {
            "id": str(i.id),
            "clause_number": i.clause_number,
            "clause_title": i.clause_title,
            "description": i.clause_description or "",
            "section": i.section or "",
            "mandatory": i.is_mandatory,
            "critical": i.is_critical,
        }
        for i in items
    ]


async def _get_evaluation_criteria(db: AsyncSession, tender_id: UUID) -> list[dict]:
    """Get evaluation criteria."""
    result = await db.execute(
        select(TenderCriteria).where(TenderCriteria.tender_id == tender_id)
    )
    criteria = result.scalars().all()
    return [
        {
            "code": c.criteria_code,
            "description": c.description or "",
            "max_marks": c.max_marks,
            "qualifying_marks": c.qualifying_marks,
            "stage": c.stage,
        }
        for c in criteria
    ]


# ---------------------------------------------------------------------------
# Main Pipeline
# ---------------------------------------------------------------------------

async def generate_solution_document(
    db: AsyncSession,
    tender_id: UUID,
    tenant_id: str,
    options: dict | None = None,
) -> dict:
    """Run the full solution document generation pipeline.

    Options:
        skip_optimization (bool): Skip the optimization pass (faster)
        sections_to_generate (list[str]): Only generate specific section numbers
        regenerate (bool): Delete existing content and regenerate

    Returns dict with generation_id and status.
    """
    options = options or {}

    # 1. Load tender
    tender_result = await db.execute(
        select(Tender).where(Tender.id == tender_id, Tender.tenant_id == UUID(tenant_id))
    )
    tender = tender_result.scalars().first()
    if not tender:
        raise ValueError("Tender not found")

    requirements_json = json.dumps(
        tender.requirements if tender.requirements else {},
        indent=2, default=str
    )[:40000]

    # 2. Create generation tracking record
    generation = ProposalGeneration(
        tender_id=tender_id,
        tenant_id=UUID(tenant_id),
        status="started",
        total_steps=8,
        completed_steps=0,
    )
    db.add(generation)
    await db.flush()
    gen_id = generation.id

    try:
        # If regenerating, clear existing proposal data
        if options.get("regenerate"):
            await db.execute(
                delete(GeneratedContent).where(GeneratedContent.tender_id == tender_id)
            )
            await db.execute(
                delete(ProposalSection).where(ProposalSection.tender_id == tender_id)
            )
            await db.execute(
                delete(TenderSection).where(TenderSection.tender_id == tender_id)
            )
            await db.flush()

        # ----- Step 1: Analyze Tender Structure -----
        generation.status = "analyzing"
        generation.current_step = "Analyzing tender structure"
        await db.flush()

        structure_prompt = STRUCTURE_ANALYSIS_PROMPT.format(
            title=tender.title or "N/A",
            reference_number=tender.reference_number or "N/A",
            issuing_authority=tender.issuing_authority or "N/A",
            requirements_json=requirements_json[:20000],
        )
        structure_text, _, _ = await _ai_call(db, tenant_id, structure_prompt, "solution_structure")
        structure_data = _parse_json_response(structure_text)

        # Save tender sections
        for i, sec in enumerate(structure_data.get("sections", [])):
            ts = TenderSection(
                tender_id=tender_id,
                section_name=sec.get("section_name", f"Section {i+1}"),
                section_description=sec.get("description"),
                section_text=json.dumps(sec.get("key_requirements", [])),
                page_reference=sec.get("page_reference"),
                sort_order=i,
            )
            db.add(ts)
        await db.flush()

        generation.completed_steps = 1
        await db.flush()

        # ----- Step 2: Classify Requirements -----
        generation.current_step = "Classifying requirements"
        await db.flush()

        requirements = await _get_requirements_list(db, tender_id)
        generation.total_requirements = len(requirements)

        # Build requirements list for prompt — use extracted JSON if no compliance items
        if requirements:
            req_text = "\n".join(
                f"  [{r['id'][:8]}] {r['clause_number']}: {r['clause_title']} "
                f"({'MANDATORY' if r['mandatory'] else 'optional'}) — {r['description'][:200]}"
                for r in requirements
            )
        else:
            req_text = f"Requirements from tender document:\n{requirements_json[:15000]}"

        classify_prompt = REQUIREMENT_CLASSIFICATION_PROMPT.format(
            title=tender.title,
            requirements_list=req_text,
        )
        classify_text, _, _ = await _ai_call(db, tenant_id, classify_prompt, "solution_classify")
        classified = _parse_json_response(classify_text)

        generation.completed_steps = 2
        await db.flush()

        # ----- Step 3: Generate TOC -----
        generation.status = "generating_toc"
        generation.current_step = "Generating table of contents"
        await db.flush()

        eval_criteria = await _get_evaluation_criteria(db, tender_id)
        eval_text = "\n".join(
            f"  - {c['code']}: {c['description']} (Max: {c['max_marks']} marks)"
            for c in eval_criteria
        ) if eval_criteria else "No evaluation criteria extracted."

        toc_prompt = TOC_GENERATION_PROMPT.format(
            title=tender.title,
            reference_number=tender.reference_number or "N/A",
            classified_sections=json.dumps(classified.get("sections", []), indent=2),
            evaluation_criteria=eval_text,
        )
        toc_text, _, _ = await _ai_call(db, tenant_id, toc_prompt, "solution_toc")
        toc_data = _parse_json_response(toc_text)

        # Save proposal sections
        toc_entries = toc_data.get("toc", [])
        section_map = {}  # section_number -> ProposalSection.id

        for i, entry in enumerate(toc_entries):
            sec_num = entry.get("section_number", str(i + 1))
            parent_num = entry.get("parent_section")
            parent_id = section_map.get(parent_num) if parent_num else None

            ps = ProposalSection(
                tender_id=tender_id,
                section_number=sec_num,
                section_title=entry.get("section_title", f"Section {sec_num}"),
                parent_section_id=parent_id,
                mapped_requirement_ids=entry.get("mapped_requirement_ids", []),
                status="pending",
                sort_order=i,
            )
            db.add(ps)
            await db.flush()
            section_map[sec_num] = ps.id

        generation.total_sections = len(toc_entries)
        generation.completed_steps = 3
        await db.flush()

        # ----- Step 4: Generate Content Per Section -----
        generation.status = "generating_sections"
        await db.flush()

        company_profile_text = await _get_company_profile_text(db, tenant_id)
        experience_text = await _get_experience_text(db, tenant_id)
        kb_content = await _get_kb_content(db, tenant_id)

        # Build requirement lookup
        req_lookup = {r["id"]: r for r in requirements}

        # Load saved proposal sections
        ps_result = await db.execute(
            select(ProposalSection)
            .where(ProposalSection.tender_id == tender_id)
            .order_by(ProposalSection.sort_order)
        )
        proposal_sections = ps_result.scalars().all()

        sections_to_gen = options.get("sections_to_generate")

        for ps in proposal_sections:
            if sections_to_gen and ps.section_number not in sections_to_gen:
                continue

            generation.current_step = f"Generating section {ps.section_number}: {ps.section_title}"
            await db.flush()

            # Gather mapped requirements text
            mapped_req_ids = ps.mapped_requirement_ids or []
            mapped_reqs_text = ""
            if mapped_req_ids and requirements:
                mapped_items = [req_lookup.get(rid) for rid in mapped_req_ids if req_lookup.get(rid)]
                if mapped_items:
                    mapped_reqs_text = "\n".join(
                        f"  - [{m['clause_number']}] {m['clause_title']}: {m['description'][:300]}"
                        for m in mapped_items
                    )
            if not mapped_reqs_text:
                mapped_reqs_text = "No specific requirements mapped to this section. Generate based on tender context."

            # Find relevant evaluation criteria
            section_eval = ""
            for ec in eval_criteria:
                if ec["description"] and (
                    ps.section_title.lower() in ec["description"].lower()
                    or ec["description"].lower() in ps.section_title.lower()
                ):
                    section_eval += f"  - {ec['code']}: {ec['description']} (Max: {ec['max_marks']})\n"
            if not section_eval:
                section_eval = "No specific evaluation criteria mapped to this section."

            # Estimate pages/words
            est_pages = 3  # default
            for entry in toc_entries:
                if entry.get("section_number") == ps.section_number:
                    est_pages = entry.get("estimated_pages", 3)
                    break
            word_count = est_pages * 350

            content_prompt = SECTION_CONTENT_PROMPT.format(
                title=tender.title,
                reference_number=tender.reference_number or "N/A",
                section_number=ps.section_number,
                section_title=ps.section_title,
                mapped_requirements=mapped_reqs_text,
                company_profile=company_profile_text,
                relevant_experience=experience_text,
                kb_content=kb_content[:3000],
                evaluation_criteria=section_eval,
                estimated_pages=est_pages,
                word_count=word_count,
            )

            content_text, tokens, model = await _ai_call(
                db, tenant_id, content_prompt, "solution_section",
                max_tokens=min(word_count * 3, 8000),
            )

            # Save generated content
            gc = GeneratedContent(
                proposal_section_id=ps.id,
                tender_id=tender_id,
                content=content_text,
                content_format="markdown",
                version=1,
                is_current=True,
                model_used=model,
                tokens_used=tokens,
            )
            db.add(gc)

            ps.status = "generated"
            await db.flush()

        generation.completed_steps = 4
        await db.flush()

        # ----- Step 5: Diagram Suggestions -----
        generation.current_step = "Suggesting diagrams"
        await db.flush()

        sections_summary = "\n".join(
            f"  {ps.section_number}. {ps.section_title}"
            for ps in proposal_sections
        )
        diagram_prompt = DIAGRAM_SUGGESTION_PROMPT.format(
            title=tender.title,
            sections_summary=sections_summary,
        )
        diagram_text, _, _ = await _ai_call(db, tenant_id, diagram_prompt, "solution_diagrams")
        try:
            diagram_data = _parse_json_response(diagram_text)
            generation.diagram_suggestions = diagram_data.get("diagrams", [])
        except Exception:
            generation.diagram_suggestions = []

        generation.completed_steps = 5
        await db.flush()

        # ----- Step 6: Gap Analysis -----
        generation.status = "gap_analysis"
        generation.current_step = "Running gap analysis"
        await db.flush()

        # Reload generated content
        gc_result = await db.execute(
            select(GeneratedContent, ProposalSection)
            .join(ProposalSection, GeneratedContent.proposal_section_id == ProposalSection.id)
            .where(
                GeneratedContent.tender_id == tender_id,
                GeneratedContent.is_current == True,
            )
            .order_by(ProposalSection.sort_order)
        )
        gc_rows = gc_result.all()

        sections_with_content = "\n\n".join(
            f"### {ps.section_number}. {ps.section_title}\n{gc.content[:2000]}"
            for gc, ps in gc_rows
        )[:30000]

        gap_prompt = GAP_ANALYSIS_PROMPT.format(
            title=tender.title,
            total_requirements=len(requirements) if requirements else "Unknown",
            requirements_list=req_text[:10000] if requirements else requirements_json[:10000],
            sections_with_content=sections_with_content,
        )
        gap_text, _, _ = await _ai_call(db, tenant_id, gap_prompt, "solution_gap_analysis")
        try:
            gap_data = _parse_json_response(gap_text)
            generation.gap_analysis = gap_data.get("gaps", [])
            generation.coverage_percentage = gap_data.get("coverage_percentage", 0)
        except Exception:
            generation.gap_analysis = []

        generation.completed_steps = 6
        await db.flush()

        # ----- Step 7: Evaluation Simulation -----
        generation.status = "evaluation"
        generation.current_step = "Simulating evaluation"
        await db.flush()

        eval_sim_prompt = EVALUATION_SIMULATION_PROMPT.format(
            title=tender.title,
            evaluation_criteria=eval_text,
            sections_with_content=sections_with_content[:25000],
            company_profile=company_profile_text,
            relevant_experience=experience_text[:3000],
        )
        eval_text_result, _, _ = await _ai_call(
            db, tenant_id, eval_sim_prompt, "solution_evaluation", max_tokens=4096
        )
        try:
            eval_data = _parse_json_response(eval_text_result)
            generation.evaluation_results = eval_data.get("evaluation", [])
            generation.estimated_score = eval_data.get("overall_percentage", 0)
        except Exception:
            generation.evaluation_results = []

        generation.completed_steps = 7
        await db.flush()

        # ----- Step 8: Optimize Weak Sections (optional) -----
        if not options.get("skip_optimization"):
            generation.status = "optimizing"
            generation.current_step = "Optimizing weak sections"
            await db.flush()

            # Find sections with low scores or gap findings
            weak_sections = set()
            if generation.evaluation_results:
                for ev in generation.evaluation_results:
                    if ev.get("percentage", 100) < 70:
                        # Find matching section
                        criteria_name = ev.get("criteria", "").lower()
                        for ps in proposal_sections:
                            if criteria_name in ps.section_title.lower() or ps.section_title.lower() in criteria_name:
                                weak_sections.add(ps.id)
                                break

            if generation.gap_analysis:
                for gap in generation.gap_analysis:
                    if gap.get("risk_level") in ("critical", "high"):
                        affected = gap.get("affected_section", "")
                        for ps in proposal_sections:
                            if ps.section_number == affected or affected in ps.section_title:
                                weak_sections.add(ps.id)
                                break

            # Optimize up to 5 weak sections
            optimized_count = 0
            for ps in proposal_sections:
                if ps.id not in weak_sections or optimized_count >= 5:
                    continue

                # Get current content
                cur_result = await db.execute(
                    select(GeneratedContent).where(
                        GeneratedContent.proposal_section_id == ps.id,
                        GeneratedContent.is_current == True,
                    )
                )
                cur_content = cur_result.scalars().first()
                if not cur_content:
                    continue

                # Gather feedback for this section
                eval_feedback = ""
                if generation.evaluation_results:
                    for ev in generation.evaluation_results:
                        if ev.get("criteria", "").lower() in ps.section_title.lower():
                            eval_feedback = json.dumps(ev, indent=2)
                            break

                gap_findings = ""
                if generation.gap_analysis:
                    relevant_gaps = [g for g in generation.gap_analysis
                                     if g.get("affected_section", "") in (ps.section_number, ps.section_title)]
                    if relevant_gaps:
                        gap_findings = json.dumps(relevant_gaps, indent=2)

                if not eval_feedback and not gap_findings:
                    continue

                mapped_req_ids = ps.mapped_requirement_ids or []
                mapped_reqs_text = ""
                if mapped_req_ids and requirements:
                    mapped_items = [req_lookup.get(rid) for rid in mapped_req_ids if req_lookup.get(rid)]
                    if mapped_items:
                        mapped_reqs_text = "\n".join(
                            f"  - [{m['clause_number']}] {m['clause_title']}"
                            for m in mapped_items
                        )

                opt_prompt = SECTION_OPTIMIZER_PROMPT.format(
                    section_number=ps.section_number,
                    section_title=ps.section_title,
                    current_content=cur_content.content[:6000],
                    evaluation_feedback=eval_feedback or "No specific feedback.",
                    gap_findings=gap_findings or "No gaps identified.",
                    mapped_requirements=mapped_reqs_text or "See original requirements.",
                )

                opt_text, tokens, model = await _ai_call(
                    db, tenant_id, opt_prompt, "solution_optimize", max_tokens=6000
                )

                # Mark old version as not current
                cur_content.is_current = False

                # Save optimized version
                new_gc = GeneratedContent(
                    proposal_section_id=ps.id,
                    tender_id=tender_id,
                    content=opt_text,
                    content_format="markdown",
                    version=cur_content.version + 1,
                    is_current=True,
                    model_used=model,
                    tokens_used=tokens,
                )
                db.add(new_gc)
                await db.flush()
                optimized_count += 1

        generation.completed_steps = 8
        generation.status = "completed"
        generation.current_step = None
        generation.completed_at = datetime.now(timezone.utc)
        await db.flush()

        return {
            "generation_id": str(gen_id),
            "status": "completed",
            "total_sections": generation.total_sections,
            "total_requirements": generation.total_requirements,
            "coverage_percentage": generation.coverage_percentage,
            "estimated_score": generation.estimated_score,
            "gap_count": len(generation.gap_analysis) if generation.gap_analysis else 0,
            "diagram_count": len(generation.diagram_suggestions) if generation.diagram_suggestions else 0,
        }

    except Exception as e:
        logger.error(f"Solution generation failed: {e}", exc_info=True)
        generation.status = "failed"
        generation.error_message = str(e)
        generation.completed_at = datetime.now(timezone.utc)
        await db.flush()
        raise


# ---------------------------------------------------------------------------
# Query functions
# ---------------------------------------------------------------------------

async def get_generation_status(db: AsyncSession, tender_id: UUID) -> dict | None:
    """Get the latest generation run status for a tender."""
    result = await db.execute(
        select(ProposalGeneration)
        .where(ProposalGeneration.tender_id == tender_id)
        .order_by(ProposalGeneration.started_at.desc())
        .limit(1)
    )
    gen = result.scalars().first()
    if not gen:
        return None

    return {
        "generation_id": str(gen.id),
        "status": gen.status,
        "current_step": gen.current_step,
        "total_steps": gen.total_steps,
        "completed_steps": gen.completed_steps,
        "progress_percentage": round(gen.completed_steps / gen.total_steps * 100) if gen.total_steps > 0 else 0,
        "total_sections": gen.total_sections,
        "total_requirements": gen.total_requirements,
        "coverage_percentage": gen.coverage_percentage,
        "estimated_score": gen.estimated_score,
        "gap_analysis": gen.gap_analysis,
        "evaluation_results": gen.evaluation_results,
        "diagram_suggestions": gen.diagram_suggestions,
        "started_at": gen.started_at.isoformat() if gen.started_at else None,
        "completed_at": gen.completed_at.isoformat() if gen.completed_at else None,
        "error_message": gen.error_message,
    }


async def get_proposal_sections(db: AsyncSession, tender_id: UUID) -> list[dict]:
    """Get all proposal sections with their current content."""
    ps_result = await db.execute(
        select(ProposalSection)
        .where(ProposalSection.tender_id == tender_id)
        .order_by(ProposalSection.sort_order)
    )
    sections = ps_result.scalars().all()

    result = []
    for ps in sections:
        # Get current content
        gc_result = await db.execute(
            select(GeneratedContent).where(
                GeneratedContent.proposal_section_id == ps.id,
                GeneratedContent.is_current == True,
            )
        )
        content = gc_result.scalars().first()

        result.append({
            "id": str(ps.id),
            "section_number": ps.section_number,
            "section_title": ps.section_title,
            "parent_section_id": str(ps.parent_section_id) if ps.parent_section_id else None,
            "mapped_requirement_ids": ps.mapped_requirement_ids or [],
            "status": ps.status,
            "sort_order": ps.sort_order,
            "content": content.content if content else None,
            "content_version": content.version if content else 0,
            "content_format": content.content_format if content else None,
            "model_used": content.model_used if content else None,
            "diagram_suggestions": content.diagram_suggestions if content else None,
            "quality_score": content.quality_score if content else None,
        })

    return result


async def regenerate_section(
    db: AsyncSession,
    tender_id: UUID,
    tenant_id: str,
    section_id: UUID,
    custom_instructions: str | None = None,
) -> dict:
    """Regenerate a single proposal section."""
    # Get the section
    ps_result = await db.execute(
        select(ProposalSection).where(
            ProposalSection.id == section_id,
            ProposalSection.tender_id == tender_id,
        )
    )
    ps = ps_result.scalars().first()
    if not ps:
        raise ValueError("Proposal section not found")

    # Load tender
    tender_result = await db.execute(select(Tender).where(Tender.id == tender_id))
    tender = tender_result.scalars().first()
    if not tender:
        raise ValueError("Tender not found")

    # Load requirements
    requirements = await _get_requirements_list(db, tender_id)
    req_lookup = {r["id"]: r for r in requirements}

    # Build mapped requirements
    mapped_req_ids = ps.mapped_requirement_ids or []
    mapped_reqs_text = ""
    if mapped_req_ids and requirements:
        mapped_items = [req_lookup.get(rid) for rid in mapped_req_ids if req_lookup.get(rid)]
        if mapped_items:
            mapped_reqs_text = "\n".join(
                f"  - [{m['clause_number']}] {m['clause_title']}: {m['description'][:300]}"
                for m in mapped_items
            )
    if not mapped_reqs_text:
        mapped_reqs_text = "Generate based on tender context."

    company_profile_text = await _get_company_profile_text(db, tenant_id)
    experience_text = await _get_experience_text(db, tenant_id)
    kb_content = await _get_kb_content(db, tenant_id)

    prompt = SECTION_CONTENT_PROMPT.format(
        title=tender.title,
        reference_number=tender.reference_number or "N/A",
        section_number=ps.section_number,
        section_title=ps.section_title,
        mapped_requirements=mapped_reqs_text,
        company_profile=company_profile_text,
        relevant_experience=experience_text,
        kb_content=kb_content[:3000],
        evaluation_criteria="Regeneration — maximize quality.",
        estimated_pages=3,
        word_count=1000,
    )

    if custom_instructions:
        prompt += f"\n\nADDITIONAL INSTRUCTIONS FROM USER:\n{custom_instructions}"

    content_text, tokens, model = await _ai_call(
        db, tenant_id, prompt, "solution_section_regen", max_tokens=6000
    )

    # Get current version number
    cur_result = await db.execute(
        select(GeneratedContent).where(
            GeneratedContent.proposal_section_id == ps.id,
            GeneratedContent.is_current == True,
        )
    )
    cur = cur_result.scalars().first()
    new_version = (cur.version + 1) if cur else 1

    if cur:
        cur.is_current = False

    gc = GeneratedContent(
        proposal_section_id=ps.id,
        tender_id=tender_id,
        content=content_text,
        content_format="markdown",
        version=new_version,
        is_current=True,
        model_used=model,
        tokens_used=tokens,
        generation_prompt=custom_instructions,
    )
    db.add(gc)
    ps.status = "generated"
    await db.flush()

    return {
        "section_id": str(ps.id),
        "section_number": ps.section_number,
        "section_title": ps.section_title,
        "content": content_text,
        "version": new_version,
        "model_used": model,
    }


async def export_proposal_docx(db: AsyncSession, tender_id: UUID, tenant_id: str) -> str:
    """Export the full proposal as a DOCX document. Returns the file path."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import os
    import tempfile

    # Load tender
    tender_result = await db.execute(select(Tender).where(Tender.id == tender_id))
    tender = tender_result.scalars().first()
    if not tender:
        raise ValueError("Tender not found")

    # Load company profile for cover page
    company_text = await _get_company_profile_text(db, tenant_id)

    # Load sections with content
    sections = await get_proposal_sections(db, tender_id)
    if not sections:
        raise ValueError("No proposal sections found. Generate the proposal first.")

    doc = DocxDocument()

    # ----- Cover Page -----
    doc.add_paragraph("")  # spacing
    doc.add_paragraph("")

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("TECHNICAL PROPOSAL")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

    doc.add_paragraph("")

    tender_title_para = doc.add_paragraph()
    tender_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tender_title_para.add_run(tender.title or "Tender Response")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph("")

    if tender.reference_number:
        ref_para = doc.add_paragraph()
        ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ref_para.add_run(f"Reference: {tender.reference_number}")
        run.font.size = Pt(12)

    if tender.issuing_authority:
        auth_para = doc.add_paragraph()
        auth_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = auth_para.add_run(f"Issuing Authority: {tender.issuing_authority}")
        run.font.size = Pt(12)

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Company info on cover
    for line in company_text.split("\n")[:5]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line).font.size = Pt(11)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f"Date: {datetime.now().strftime('%d %B %Y')}").font.size = Pt(11)

    doc.add_page_break()

    # ----- Table of Contents -----
    toc_heading = doc.add_heading("Table of Contents", level=1)

    for sec in sections:
        depth = sec["section_number"].count(".")
        indent = "    " * depth
        toc_line = f"{indent}{sec['section_number']}  {sec['section_title']}"
        p = doc.add_paragraph(toc_line)
        p.style = doc.styles["Normal"]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ----- Sections -----
    for sec in sections:
        depth = sec["section_number"].count(".")
        heading_level = min(depth + 1, 4)  # h1 to h4

        doc.add_heading(
            f"{sec['section_number']}  {sec['section_title']}",
            level=heading_level,
        )

        content = sec.get("content")
        if content:
            # Convert markdown to paragraphs (simple conversion)
            _render_markdown_to_docx(doc, content)
        else:
            p = doc.add_paragraph("[Content not yet generated]")
            p.runs[0].italic = True

    # Save to temp file
    temp_dir = tempfile.mkdtemp()
    safe_title = "".join(c for c in (tender.title or "proposal")[:50] if c.isalnum() or c in " -_")
    filename = f"Technical_Proposal_{safe_title}.docx"
    filepath = os.path.join(temp_dir, filename)
    doc.save(filepath)

    return filepath


def _render_markdown_to_docx(doc, markdown_text: str):
    """Simple markdown to DOCX renderer — handles headings, bullets, bold, tables."""
    from docx.shared import Pt

    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)

        # Bullet points
        elif line.startswith("- ") or line.startswith("* "):
            text = line[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_runs(p, text)

        # Numbered lists
        elif len(line) > 2 and line[0].isdigit() and line[1] in ".)" and line[2] == " ":
            text = line[3:]
            p = doc.add_paragraph(style="List Number")
            _add_formatted_runs(p, text)

        # Table detection (markdown pipe tables)
        elif line.startswith("|") and "|" in line[1:]:
            # Collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1  # will be incremented at end of loop
            _render_table(doc, table_lines)

        # Placeholder markers
        elif line.startswith("[Insert") or line.startswith("[Bidder"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.italic = True
            run.font.size = Pt(10)

        # Normal paragraph
        else:
            p = doc.add_paragraph()
            _add_formatted_runs(p, line)

        i += 1


def _add_formatted_runs(paragraph, text: str):
    """Add text with bold/italic formatting from markdown."""
    import re
    # Split by bold markers **text**
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _render_table(doc, table_lines: list[str]):
    """Render a markdown pipe table into a DOCX table."""
    from docx.shared import Pt

    rows_data = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        # Skip separator rows (---|----|---)
        if all(set(c.strip()) <= {"-", ":"} for c in cells):
            continue
        rows_data.append(cells)

    if not rows_data:
        return

    num_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(rows_data):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.cell(row_idx, col_idx)
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    # Bold the header row
    if rows_data:
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

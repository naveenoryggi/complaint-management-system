"""DOCX document formatting service with letterhead and signature support."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from io import BytesIO
from typing import Optional, Dict, Any
import os
from datetime import datetime

from app.core.config import settings


class DOCXService:
    """Service for creating formatted DOCX documents with letterhead support."""

    def __init__(self):
        """Initialize DOCX service."""
        self.default_font = "Calibri"
        self.heading_font = "Arial"

    def create_document_from_text(
        self,
        content: str,
        title: str,
        generation_type: str,
        company_name: str = None,
        add_letterhead: bool = True,
        company_profile: Dict[str, Any] = None,
        add_signature: bool = False,
        signatory_name: str = None,
        designation: str = None,
    ) -> BytesIO:
        """
        Create a formatted DOCX document from text content.

        If a letterhead DOCX template is uploaded, uses it as the base document
        (preserving its headers, footers, and page setup). Otherwise falls back
        to programmatic letterhead generation.

        Args:
            content: Text content
            title: Document title
            generation_type: Type of document (technical_solution, etc.)
            company_name: Company name for simple letterhead
            add_letterhead: Whether to add header/footer
            company_profile: Full company profile dict for enhanced letterhead
            add_signature: Whether to add signature block at the end
            signatory_name: Name of authorized signatory
            designation: Signatory designation

        Returns:
            BytesIO object containing the DOCX file
        """
        # Try to use uploaded letterhead DOCX as base template
        doc = self._load_letterhead_template(company_profile) if add_letterhead else None
        using_template = doc is not None

        if doc:
            # Letterhead template loaded — styles/headers/footers come from template
            self._configure_styles(doc)
        else:
            # No template — create blank document with programmatic letterhead
            doc = Document()
            self._configure_styles(doc)

            if add_letterhead:
                if company_profile:
                    self._add_enhanced_letterhead(doc, company_profile)
                else:
                    self._add_letterhead(doc, company_name or "Your Company")

        # Add document title
        if generation_type != "covering_letter":
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.font.size = Pt(12)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(0, 51, 102)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

        # Parse and format content
        self._parse_and_format_content(doc, content, generation_type)

        # Add signature in footer (appears on every page)
        if add_signature:
            effective_name = company_profile.get("company_name", company_name or "Your Company") if company_profile else (company_name or "Your Company")
            self._add_signature_footer(doc, effective_name, company_profile, signatory_name, designation)
        elif add_letterhead and not using_template:
            # No signature requested — add a simple footer instead
            footer_name = company_profile.get("company_name", company_name or "Your Company") if company_profile else (company_name or "Your Company")
            footer_address = company_profile.get("registered_address") if company_profile else None
            self._add_enhanced_footer(doc, footer_name, footer_address)

        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _configure_styles(self, doc: Document):
        """Configure document styles."""
        styles = doc.styles

        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = self.default_font
        normal_font.size = Pt(11)

        try:
            heading1 = styles['Heading 1']
            heading1.font.name = self.heading_font
            heading1.font.size = Pt(12)
            heading1.font.bold = True
            heading1.font.color.rgb = RGBColor(0, 51, 102)
        except KeyError:
            pass

        try:
            heading2 = styles['Heading 2']
            heading2.font.name = self.heading_font
            heading2.font.size = Pt(12)
            heading2.font.bold = True
            heading2.font.color.rgb = RGBColor(51, 102, 153)
        except KeyError:
            pass

    def _load_letterhead_template(self, company_profile: Dict[str, Any] = None) -> Optional[Document]:
        """
        Load the uploaded letterhead DOCX as the base document template.

        Clears the body content but preserves headers, footers, page setup,
        and any watermarks/background images defined in the template.

        Returns:
            Document if letterhead template found and loaded, else None
        """
        if not company_profile:
            return None

        letterhead_path = company_profile.get("letterhead_path")
        if not letterhead_path:
            return None

        # Only use DOCX letterhead templates (not images)
        if not letterhead_path.lower().endswith(('.docx', '.doc')):
            return None

        full_path = os.path.join(settings.upload_dir, letterhead_path)
        if not os.path.exists(full_path):
            return None

        try:
            doc = Document(full_path)

            # Clear existing body content from the template, keeping
            # headers/footers/sections intact
            body = doc.element.body
            for child in list(body):
                # Keep sectPr (section properties — contains header/footer refs)
                if child.tag.endswith('}sectPr'):
                    continue
                body.remove(child)

            return doc
        except Exception:
            return None

    def _add_letterhead(self, doc: Document, company_name: str):
        """Add simple company letterhead to document header."""
        section = doc.sections[0]
        header = section.header

        header_para = header.paragraphs[0]
        header_run = header_para.add_run(company_name)
        header_run.font.size = Pt(14)
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(0, 51, 102)
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        header.add_paragraph("_" * 80)

    def _add_enhanced_letterhead(self, doc: Document, company_profile: Dict[str, Any]):
        """
        Add enhanced letterhead with logo, company details, and separator.

        Args:
            doc: Document object
            company_profile: Dict with company_name, logo_path, registered_address,
                           phone, email, website, pan_number, gstin, etc.
        """
        section = doc.sections[0]
        # Set margins
        section.top_margin = Cm(1.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        header = section.header

        company_name = company_profile.get("company_name", "Your Company")
        logo_path = company_profile.get("logo_path")

        # Try to embed logo
        logo_added = False
        if logo_path:
            full_logo_path = os.path.join(settings.upload_dir, logo_path)
            if os.path.exists(full_logo_path):
                try:
                    logo_para = header.paragraphs[0]
                    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    logo_run = logo_para.add_run()
                    logo_run.add_picture(full_logo_path, height=Cm(2))
                    logo_added = True
                except Exception:
                    pass

        # Company name
        if logo_added:
            name_para = header.add_paragraph()
        else:
            name_para = header.paragraphs[0]
        name_run = name_para.add_run(company_name)
        name_run.font.size = Pt(14)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0, 51, 102)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Address line
        address = company_profile.get("registered_address")
        if address:
            addr_para = header.add_paragraph()
            addr_run = addr_para.add_run(address)
            addr_run.font.size = Pt(9)
            addr_run.font.color.rgb = RGBColor(80, 80, 80)
            addr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact line: Phone | Email | Website
        contact_parts = []
        if company_profile.get("phone"):
            contact_parts.append(f"Phone: {company_profile['phone']}")
        if company_profile.get("email"):
            contact_parts.append(f"Email: {company_profile['email']}")
        if company_profile.get("website"):
            contact_parts.append(company_profile["website"])

        if contact_parts:
            contact_para = header.add_paragraph()
            contact_run = contact_para.add_run(" | ".join(contact_parts))
            contact_run.font.size = Pt(8)
            contact_run.font.color.rgb = RGBColor(100, 100, 100)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # PAN / GSTIN line
        id_parts = []
        if company_profile.get("pan_number"):
            id_parts.append(f"PAN: {company_profile['pan_number']}")
        if company_profile.get("gstin"):
            id_parts.append(f"GSTIN: {company_profile['gstin']}")
        if company_profile.get("cin_number"):
            id_parts.append(f"CIN: {company_profile['cin_number']}")

        if id_parts:
            id_para = header.add_paragraph()
            id_run = id_para.add_run(" | ".join(id_parts))
            id_run.font.size = Pt(8)
            id_run.font.color.rgb = RGBColor(100, 100, 100)
            id_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Horizontal separator
        sep_para = header.add_paragraph()
        sep_run = sep_para.add_run("_" * 85)
        sep_run.font.size = Pt(8)
        sep_run.font.color.rgb = RGBColor(0, 51, 102)
        sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_enhanced_footer(self, doc: Document, company_name: str, address: str = None):
        """Add enhanced footer with company address and generation date."""
        section = doc.sections[0]
        footer = section.footer

        # Separator
        sep_para = footer.paragraphs[0]
        sep_run = sep_para.add_run("_" * 85)
        sep_run.font.size = Pt(8)
        sep_run.font.color.rgb = RGBColor(0, 51, 102)
        sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Company info + date
        parts = [company_name]
        if address:
            # Take first line of address for footer
            first_line = address.split("\n")[0].strip()
            if first_line:
                parts.append(first_line)

        footer_text = " | ".join(parts)
        footer_text += f" | Generated on {datetime.utcnow().strftime('%B %d, %Y')}"

        info_para = footer.add_paragraph()
        info_run = info_para.add_run(footer_text)
        info_run.font.size = Pt(8)
        info_run.font.color.rgb = RGBColor(128, 128, 128)
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_footer(self, doc: Document, company_name: str):
        """Add simple footer (backwards compatible)."""
        self._add_enhanced_footer(doc, company_name)

    def _add_signature_footer(
        self,
        doc: Document,
        company_name: str,
        company_profile: Dict[str, Any] = None,
        signatory_name: str = None,
        designation: str = None,
    ):
        """
        Add signature block in the document footer so it appears on EVERY page.

        Includes signature image, stamp image, signatory name, and designation.
        """
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False

        # Separator line
        sep_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        sep_run = sep_para.add_run("_" * 85)
        sep_run.font.size = Pt(7)
        sep_run.font.color.rgb = RGBColor(180, 180, 180)
        sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Signature + stamp images side by side (if available)
        sig_added = False
        if company_profile and company_profile.get("signature_path"):
            sig_path = os.path.join(settings.upload_dir, company_profile["signature_path"])
            if os.path.exists(sig_path):
                try:
                    sig_para = footer.add_paragraph()
                    sig_run = sig_para.add_run()
                    sig_run.add_picture(sig_path, height=Cm(1.2))
                    sig_added = True
                except Exception:
                    pass

        if company_profile and company_profile.get("stamp_path"):
            stamp_path = os.path.join(settings.upload_dir, company_profile["stamp_path"])
            if os.path.exists(stamp_path):
                try:
                    # Add stamp to same paragraph as signature if possible
                    if sig_added:
                        sig_para = footer.paragraphs[-1]
                        stamp_run = sig_para.add_run("   ")
                        stamp_run = sig_para.add_run()
                        stamp_run.add_picture(stamp_path, height=Cm(1.2))
                    else:
                        stamp_para = footer.add_paragraph()
                        stamp_run = stamp_para.add_run()
                        stamp_run.add_picture(stamp_path, height=Cm(1.2))
                except Exception:
                    pass

        # Signatory line: "For <company> | Name | Designation"
        sig_parts = [f"For and on behalf of {company_name}"]
        if signatory_name:
            sig_parts.append(signatory_name)
        if designation:
            sig_parts.append(designation)

        sig_text_para = footer.add_paragraph()
        sig_text_run = sig_text_para.add_run(" | ".join(sig_parts))
        sig_text_run.font.size = Pt(8)
        sig_text_run.font.bold = True
        sig_text_run.font.color.rgb = RGBColor(30, 30, 30)
        sig_text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Date
        date_para = footer.add_paragraph()
        date_run = date_para.add_run(f"Date: {datetime.utcnow().strftime('%d/%m/%Y')}")
        date_run.font.size = Pt(8)
        date_run.font.color.rgb = RGBColor(100, 100, 100)
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _has_style(self, doc: Document, style_name: str) -> bool:
        """Check if a style exists in the document."""
        try:
            doc.styles[style_name]
            return True
        except KeyError:
            return False

    def _parse_and_format_content(
        self,
        doc: Document,
        content: str,
        generation_type: str
    ):
        """Parse text content and format appropriately."""
        lines = content.split('\n')

        has_bullet = self._has_style(doc, 'List Bullet')
        has_number = self._has_style(doc, 'List Number')

        is_declaration = generation_type == "declaration"
        found_first_heading = False

        for line in lines:
            line = line.strip()

            if not line:
                doc.add_paragraph()
                continue

            # For declarations: only the FIRST all-caps line is a heading (title).
            # Subsequent all-caps lines (sub-sections) are bold paragraphs, not headings.
            if is_declaration:
                is_h1 = self._is_heading_1(line)
                if is_h1 and not found_first_heading:
                    found_first_heading = True
                    para = doc.add_paragraph()
                    run = para.add_run(line.lstrip('#').strip())
                    run.font.bold = True
                    run.font.size = Pt(12)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif is_h1:
                    # Sub-section heading: bold but normal font size
                    para = doc.add_paragraph()
                    run = para.add_run(line)
                    run.font.bold = True
                    run.font.size = Pt(11)
                elif line.startswith('- ') or line.startswith('* '):
                    if has_bullet:
                        para = doc.add_paragraph(line[2:], style='List Bullet')
                    else:
                        para = doc.add_paragraph(f"•  {line[2:]}")
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                elif line[0].isdigit() and '. ' in line[:5]:
                    para = doc.add_paragraph(line)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    para = doc.add_paragraph(line)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                # Non-declaration content: original heading detection
                if self._is_heading_1(line):
                    para = doc.add_heading(line.lstrip('#').strip(), level=1)
                elif self._is_heading_2(line):
                    para = doc.add_heading(line.lstrip('#').strip(), level=2)
                elif self._is_heading_3(line):
                    para = doc.add_heading(line.lstrip('#').strip(), level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    if has_bullet:
                        para = doc.add_paragraph(line[2:], style='List Bullet')
                    else:
                        para = doc.add_paragraph(f"•  {line[2:]}")
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                elif line[0].isdigit() and '. ' in line[:5]:
                    if has_number:
                        para = doc.add_paragraph(line.split('. ', 1)[1], style='List Number')
                    else:
                        para = doc.add_paragraph(line)
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    para = doc.add_paragraph(line)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _is_heading_1(self, line: str) -> bool:
        """Check if line is a level 1 heading."""
        if line.startswith('# ') and not line.startswith('## '):
            return True
        if line.isupper() and len(line) < 60 and not line.endswith('.'):
            return True
        heading_keywords = [
            'TECHNICAL SOLUTION',
            'EXECUTIVE SUMMARY',
            'DECLARATION',
            'METHODOLOGY',
            'INTRODUCTION',
            'CONCLUSION',
        ]
        if any(keyword in line.upper() for keyword in heading_keywords):
            return True
        return False

    def _is_heading_2(self, line: str) -> bool:
        """Check if line is a level 2 heading."""
        if line.startswith('## ') and not line.startswith('### '):
            return True
        if line[0].isdigit() and '. ' in line[:10] and len(line.split('. ', 1)[1].split()) < 8:
            return True
        return False

    def _is_heading_3(self, line: str) -> bool:
        """Check if line is a level 3 heading."""
        if line.startswith('### '):
            return True
        return False

    def format_covering_letter(
        self,
        content: str,
        company_name: str,
        address: Optional[str] = None,
    ) -> BytesIO:
        """Format a covering letter with proper business letter layout."""
        doc = Document()
        self._configure_styles(doc)

        header_para = doc.add_paragraph()
        header_run = header_para.add_run(company_name)
        header_run.font.size = Pt(16)
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(0, 51, 102)
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if address:
            addr_para = doc.add_paragraph(address)
            addr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            addr_para.runs[0].font.size = Pt(10)

        doc.add_paragraph()
        doc.add_paragraph("_" * 80)
        doc.add_paragraph()

        self._parse_and_format_content(doc, content, "covering_letter")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def format_compliance_declaration(
        self,
        content: str,
        company_name: str,
    ) -> BytesIO:
        """Format a compliance declaration with signature section."""
        doc = Document()
        self._configure_styles(doc)

        title = doc.add_heading("DECLARATION", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        self._parse_and_format_content(doc, content, "compliance_declaration")

        doc.add_paragraph()
        doc.add_paragraph()

        sig_table = doc.add_table(rows=5, cols=2)
        sig_table.style = 'Table Grid'

        sig_table.rows[0].cells[0].text = "Authorized Signatory:"
        sig_table.rows[1].cells[0].text = "Name:"
        sig_table.rows[2].cells[0].text = "Designation:"
        sig_table.rows[3].cells[0].text = "Date:"
        sig_table.rows[4].cells[0].text = "Company Stamp:"

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer


# Global service instance
docx_service = DOCXService()
